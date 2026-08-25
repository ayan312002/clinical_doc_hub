"""Generate a linked synthetic patient-document dataset for testing.

Produces ~35 documents across 9 patients (shared MRNs) in samples/:
- Formats: PDF (reportlab), DOCX (python-docx), TXT, PNG images (OCR path)
- Acuity mix: routine / high / critical to exercise the triage rules engine

Run: python scripts/generate_dataset.py
"""

import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

SAMPLES_DIR = Path(__file__).parent.parent / "samples"
SAMPLES_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Writers
# ---------------------------------------------------------------------------

def _render_lines(text: str):
    """Yield (line, is_heading) pairs from a plain-text document body."""
    for raw in text.strip().split("\n"):
        line = raw.strip()
        yield line, bool(line) and line.isupper() and len(line) > 3


def create_pdf(text: str, filename: str):
    filepath = SAMPLES_DIR / filename
    doc = SimpleDocTemplate(str(filepath), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    for line, is_heading in _render_lines(text):
        if not line:
            story.append(Spacer(1, 6))
        elif is_heading:
            story.append(Paragraph(f"<b>{line}</b>", styles["Heading2"]))
        else:
            story.append(Paragraph(line, styles["Normal"]))
    doc.build(story)
    print(f"Created: {filepath.name}")


def create_docx(text: str, filename: str):
    filepath = SAMPLES_DIR / filename
    d = DocxDocument()
    for line, is_heading in _render_lines(text):
        if not line:
            continue
        if is_heading:
            run = d.add_paragraph().add_run(line)
            run.bold = True
        else:
            d.add_paragraph(line)
    d.save(str(filepath))
    print(f"Created: {filepath.name}")


def create_txt(text: str, filename: str):
    filepath = SAMPLES_DIR / filename
    filepath.write_text(text.strip())
    print(f"Created: {filepath.name}")


def _load_fonts():
    try:
        body = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf", 16
        )
        heading = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf", 18
        )
    except OSError:
        body = heading = ImageFont.load_default()
    return body, heading


def create_png(text: str, filename: str):
    filepath = SAMPLES_DIR / filename
    body_font, heading_font = _load_fonts()

    wrapped = []
    for line, is_heading in _render_lines(text):
        segments = textwrap.wrap(line, width=95) if line else [""]
        for segment in segments:
            wrapped.append((segment, is_heading))

    line_height = 26
    margin = 40
    width = 1300
    height = margin * 2 + line_height * len(wrapped)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    y = margin
    for segment, is_heading in wrapped:
        font = heading_font if is_heading else body_font
        draw.text((margin, y), segment, fill="black", font=font)
        y += line_height
    img.save(str(filepath))
    print(f"Created: {filepath.name}")


# ---------------------------------------------------------------------------
# Patients 1-3: existing patients (fill missing documents, extend timelines)
# ---------------------------------------------------------------------------

INTAKE_THOMPSON = """
St. Mary's Medical Center
PATIENT INTAKE FORM (EMERGENCY DEPARTMENT)

Date of Visit: 01/05/2026
Arrival Time: 21:47 via EMS

PATIENT INFORMATION:
First Name: Margaret
Last Name: Thompson
Date of Birth: 03/15/1958 (Age 67)
Sex: Female
MRN: MRN-2024-0847
Phone: (555) 118-2044
Address: 88 Chestnut Lane, Springfield, IL 62702
Insurance: Medicare
Policy Number: MCR-445-8821

EMERGENCY CONTACT:
Name: Ellen Thompson (Daughter)
Phone: (555) 118-9087

CHIEF COMPLAINT:
Progressive shortness of breath and leg swelling for 3 days

HISTORY OF PRESENT ILLNESS:
68-year-old female brought in by EMS with 3 days of worsening dyspnea on exertion,
orthopnea (now sleeping upright in a recliner), bilateral lower extremity swelling,
and an 8 lb weight gain over the past week. Reports non-compliance with low-sodium
diet over the holidays. No chest pain, fever, or cough with purulent sputum.

PAST MEDICAL HISTORY:
1. Chronic heart failure (HFrEF, EF 35%, documented 2023)
2. Type 2 Diabetes Mellitus
3. Hypertension
4. Chronic kidney disease stage 2

CURRENT MEDICATIONS (PER HOME LIST):
1. Lisinopril 20 mg PO daily
2. Furosemide 40 mg PO daily
3. Metformin 1000 mg PO BID
4. Carvedilol 12.5 mg PO BID
5. Aspirin 81 mg PO daily

ALLERGIES:
- Penicillin (rash)
- Sulfa drugs (hives)

SOCIAL HISTORY:
- Retired schoolteacher, lives with daughter
- Tobacco: Former smoker (20 pack-years, quit 2005)
- Alcohol: None

VITAL SIGNS ON ARRIVAL:
- Blood Pressure: 168/94 mmHg
- Heart Rate: 102 bpm
- Temperature: 98.6 F
- Respiratory Rate: 24 breaths/min
- O2 Saturation: 89% on room air
- Weight: 168 lbs

ED COURSE INITIALLY NOTED:
Placed on 2L nasal cannula, IV access obtained, chest x-ray and stat labs ordered.
Cardiology consulted for suspected acute decompensated heart failure.
"""

LAB_THOMPSON = """
St. Mary's Medical Center
LABORATORY REPORT

Patient: Margaret Thompson
DOB: 03/15/1958 (Age 67)
MRN: MRN-2024-0847
Ordering Physician: Dr. James Chen, MD
Collection Date: 01/06/2026 (01:12)

CARDIAC MARKERS:
- BNP: 1200 pg/mL (HIGH) [Ref: <100]
- Troponin I: 0.02 ng/mL (Normal) [Ref: <0.04]
- CK-MB: 2.1 ng/mL (Normal) [Ref: 0.0-5.0]

COMPLETE BLOOD COUNT (CBC):
- White Blood Cell Count: 9.8 x10^3/uL (Normal) [Ref: 4.5-11.0]
- Hemoglobin: 10.9 g/dL (LOW) [Ref: 12.0-15.5]
- Hematocrit: 33.8% (LOW) [Ref: 34.9-44.5]
- Platelet Count: 287 x10^3/uL (Normal) [Ref: 150-400]

COMPREHENSIVE METABOLIC PANEL (CMP):
- Glucose: 176 mg/dL (HIGH) [Ref: 70-100]
- BUN: 34 mg/dL (HIGH) [Ref: 7-20]
- Creatinine: 1.4 mg/dL (HIGH) [Ref: 0.6-1.1]
- eGFR: 42 mL/min/1.73m2 (LOW) [Ref: >60]
- Sodium: 136 mEq/L (Normal) [Ref: 136-145]
- Potassium: 5.1 mEq/L (Normal) [Ref: 3.5-5.0]
- Chloride: 99 mEq/L (Normal) [Ref: 98-106]
- CO2: 27 mEq/L (Normal) [Ref: 23-29]
- Calcium: 8.9 mg/dL (Normal) [Ref: 8.5-10.5]

MAGNESIUM: 1.9 mg/dL (Normal) [Ref: 1.7-2.2]

INTERPRETATION:
Findings consistent with acutely decompensated heart failure: markedly raised BNP,
prerenal azotemia with reduced eGFR, mild anemia of chronic disease, and poor
glycemic control. Recommend serial renal function monitoring during diuresis.

Reported by: Core Laboratory
Verified by: Dr. Michael Torres, MD, Pathologist
Date: 01/06/2026
"""

INTAKE_MARTINEZ = """
Valley Health Clinic
NEW PATIENT INTAKE FORM

Date of Visit: 01/09/2026

PATIENT INFORMATION:
First Name: Robert
Last Name: Martinez
Date of Birth: 07/22/1975 (Age 50)
Sex: Male
MRN: MRN-2024-1203
Phone: (555) 771-3320
Email: r.martinez@email.com
Address: 1209 Maple Street, Springfield, IL 62703
Insurance: Aetna PPO
Policy Number: AET-220-77109

EMERGENCY CONTACT:
Name: Lucia Martinez (Wife)
Phone: (555) 771-4415

CHIEF COMPLAINT:
Establish care; increased thirst and frequent urination for several months

HISTORY OF PRESENT ILLNESS:
50-year-old male presenting to establish primary care. Reports polyuria and
polydipsia for ~4 months, unintentional 12 lb weight loss, and blurred vision at
end of day. Father diagnosed with type 2 diabetes at age 55. No prior primary
care visits in past 5 years.

PAST MEDICAL HISTORY:
- No known chronic conditions previously diagnosed
- Right knee arthroscopy (2010)

CURRENT MEDICATIONS:
-None reported-

ALLERGIES:
- No known drug allergies

FAMILY HISTORY:
- Father: Type 2 Diabetes, died age 68 (MI)
- Mother: Obesity, Osteoarthritis
- Brother: Hyperlipidemia

SOCIAL HISTORY:
- Occupation: Warehouse supervisor
- Tobacco: Never smoker
- Alcohol: 4-6 beers per week
- Exercise: Rare
- Diet: Fast food 4-5 times per week, regular soda daily

REVIEW OF SYSTEMS:
- Constitutional: 12 lb weight loss over 4 months
- HEENT: Intermittent blurred vision
- Genitourinary: Nocturia 2-3 times nightly
- Extremities: No numbness or tingling

VITAL SIGNS:
- Blood Pressure: 142/88 mmHg
- Heart Rate: 84 bpm
- Temperature: 98.1 F
- Respiratory Rate: 16 breaths/min
- O2 Saturation: 98%
- Height: 5'11" (180.3 cm)
- Weight: 232 lbs (105.2 kg)
- BMI: 32.3 kg/m2

ASSESSMENT:
1. Suspected new-onset type 2 diabetes - diagnostic labs ordered
2. Stage 1 hypertension by office readings
3. Obesity class I
4. Dyslipidemia risk - screening ordered

PLAN:
1. Order CMP, HbA1c, lipid panel, CBC, TSH (collected 01/10)
2. Urine microalbumin screening
3. Lifestyle counseling provided today
4. Return visit 01/13 for lab review and treatment planning
"""

NOTES_MARTINEZ = """
Valley Health Clinic
PROGRESS NOTE

Date: 01/13/2026
Patient: Robert Martinez
MRN: MRN-2024-1203
Provider: Dr. Sarah Williams, MD

SUBJECTIVE:
Returns to review laboratory results obtained 01/10. Confirms ongoing polyuria and
nocturia. Has cut regular soda to one per day since last visit. Denies polyphagia,
numbness, or visual changes beyond end-of-day blurring.

OBJECTIVE:
Vital Signs:
- BP: 138/86 mmHg
- HR: 80 bpm
- Weight: 231 lbs

Lab Review (01/10):
- HbA1c: 8.2% (elevated)
- Fasting glucose: 185 mg/dL (elevated)
- Total cholesterol: 245 mg/dL (elevated); LDL 165; HDL 38; TG 210
- WBC 12.5 x10^3/uL - mild leukocytosis; urine culture sent, pending
- TSH: 2.8 mIU/L (normal)
- Urine microalbumin/creatinine ratio: 42 mcg/mg (elevated)

ASSESSMENT:
1. Type 2 diabetes mellitus, newly diagnosed, with early nephropathy signal
2. Dyslipidemia
3. Stage 1 hypertension
4. Mild leukocytosis - possible early urinary tract infection, culture pending

PLAN:
1. Start metformin 500 mg PO BID, titrate over 4 weeks
2. Start atorvastatin 20 mg PO nightly
3. Repeat urinalysis/culture follow-up call in 48 hours
4. Diabetes self-management education referral placed
5. Recheck BMP and A1c in 3 months; annual eye exam referral
6. Home BP log requested

Patient verbalized understanding and agreed with plan.
"""

INTAKE_JOHNSON = """
Valley Health Clinic
PATIENT INTAKE FORM

Date of Visit: 01/12/2026

PATIENT INFORMATION:
First Name: Emily
Last Name: Johnson
Date of Birth: 11/03/1990 (Age 35)
Sex: Female
MRN: MRN-2024-1567
Phone: (555) 234-5678
Email: emily.johnson@email.com
Address: 456 Oak Avenue, Springfield, IL 62704
Insurance: BlueCross BlueShield
Policy Number: BC-9876543

EMERGENCY CONTACT:
Name: David Johnson (Spouse)
Phone: (555) 234-5679

CHIEF COMPLAINT:
Persistent fatigue and headaches for the past 2 weeks

HISTORY OF PRESENT ILLNESS:
Patient reports worsening fatigue over the past 2 weeks, accompanied by daily
headaches (mostly frontal, mild to moderate intensity). Denies fever, chills,
vision changes, or recent trauma. Reports increased stress at work. Sleep pattern
disrupted - averaging 5-6 hours per night.

PAST MEDICAL HISTORY:
1. Migraine without aura (diagnosed 2018)
2. Generalized anxiety disorder (diagnosed 2020)
3. Iron deficiency anemia (resolved 2023)

SURGICAL HISTORY:
- Appendectomy (2015)
- Wisdom teeth extraction (2012)

CURRENT MEDICATIONS:
1. Sumatriptan 50 mg PO as needed for migraines
2. Sertraline 50 mg PO daily
3. Iron supplement 325 mg PO daily
4. Multivitamin 1 tablet PO daily

ALLERGIES:
- Codeine (nausea/vomiting)

FAMILY HISTORY:
- Mother: Hypertension, Type 2 Diabetes
- Father: Coronary artery disease (MI at age 62)
- Sister: Asthma

SOCIAL HISTORY:
- Occupation: Marketing Manager
- Tobacco: Never smoker
- Alcohol: Social, 2-3 drinks per week
- Exercise: Walking 3x/week
- Diet: Regular diet; caffeine 4-5 cups coffee daily

REVIEW OF SYSTEMS:
- Constitutional: Fatigue, no fever or weight changes
- HEENT: Headaches (frontal), no vision changes
- Cardiovascular: No chest pain or palpitations
- Respiratory: No shortness of breath
- GI: No nausea, vomiting, or diarrhea
- Neurological: Headaches, no numbness or weakness
- Psychiatric: Anxiety stable on current medication

VITAL SIGNS:
- Blood Pressure: 118/72 mmHg
- Heart Rate: 76 bpm
- Temperature: 98.4 F
- Respiratory Rate: 16 breaths/min
- O2 Saturation: 99%
- Height: 5'6" (167.6 cm)
- Weight: 135 lbs (61.2 kg)
- BMI: 21.8 kg/m2

ASSESSMENT:
1. Fatigue, likely multifactorial (stress, sleep insufficiency, possible anemia recurrence)
2. Tension-type headaches vs migraine progression
3. Generalized anxiety disorder, stable

PLAN:
1. Order CBC with iron studies, TSH, CMP
2. Sleep hygiene counseling
3. Continue current medications
4. Follow up in 2 weeks with lab results
5. Consider neurology referral if headaches worsen
"""

FOLLOWUP_NOTE_JOHNSON = """
Valley Health Clinic
FOLLOW-UP VISIT NOTE

Date: 01/24/2026
Patient: Emily Johnson
MRN: MRN-2024-1567
Provider: Dr. Sarah Williams, MD

SUBJECTIVE:
Follow-up from 01/12 visit. Returns with lab results. Reports improved energy
levels since last visit. Headaches have decreased in frequency (now 2-3 per week
vs daily). Started sleep hygiene practices - now averaging 7 hours per night.
Reduced coffee to 2 cups daily.

OBJECTIVE:
Vital Signs:
- BP: 115/70 mmHg
- HR: 72 bpm
- Weight: 134 lbs

General: Well-appearing, no acute distress

Lab Results Review:
- CBC: WBC 7.2 (normal), Hgb 12.8 (improved from previous 10.1), Hct 38.2%
- Iron studies: Ferritin 45 ng/mL (improved from 12), TIBC 350 mcg/dL, Iron sat 22%
- TSH: 2.1 mIU/L (normal)
- CMP: All within normal limits
- Glucose: 88 mg/dL (normal)

ASSESSMENT:
1. Iron deficiency anemia - resolved with supplementation
2. Tension headaches - improved with sleep hygiene and stress management
3. Anxiety disorder - stable on sertraline

PLAN:
1. Discontinue iron supplement (iron stores replete)
2. Continue sertraline 50 mg daily
3. Continue sleep hygiene practices
4. Headache diary recommended
5. Return in 3 months for routine follow-up
6. Neurology referral if headache pattern changes

Patient verbalized understanding. Questions answered.
"""

# ---------------------------------------------------------------------------
# Patient 4: Harold Brooks - 78M - urosepsis (CRITICAL acuity)
# ---------------------------------------------------------------------------

INTAKE_BROOKS = """
St. Mary's Medical Center
PATIENT INTAKE FORM (EMERGENCY DEPARTMENT)

Date of Visit: 07/28/2026
Arrival Time: 08:15 via EMS

PATIENT INFORMATION:
First Name: Harold
Last Name: Brooks
Date of Birth: 04/30/1948 (Age 78)
Sex: Male
MRN: MRN-2025-0112
Phone: (555) 402-6611
Address: 17 Prairie View Court, Springfield, IL 62711
Insurance: Medicare + Aetna Supplement
Policy Number: MCR-778-1042

EMERGENCY CONTACT:
Name: Nancy Brooks (Daughter)
Phone: (555) 402-7788

CHIEF COMPLAINT:
Fever, confusion, and decreased urination for 2 days

HISTORY OF PRESENT ILLNESS:
78-year-old male brought by family with 2 days of fever to 103.1 F at home, new
confusion (non-verbal overnight per daughter), burning urination with foul-smelling
urine, and markedly reduced urine output. History of recurrent UTIs; most recent
course completed 6 weeks ago. Found collapsed attempting to stand this morning.

PAST MEDICAL HISTORY:
1. Benign prostatic hyperplasia
2. Recurrent urinary tract infections (3 in past year)
3. Chronic kidney disease stage 3a
4. Atrial fibrillation on apixaban
5. Hearing impairment (bilateral aids)

SURGICAL HISTORY:
- TURP (2019)
- Inguinal hernia repair (2014)

HOME MEDICATIONS:
1. Tamsulosin 0.4 mg PO daily
2. Apixaban 5 mg PO BID
3. Atorvastatin 40 mg PO daily
4. Nitrofurantoin 100 mg PO BID (suppressive)

ALLERGIES:
- Codeine (hives)

FAMILY HISTORY:
- Non-contributory per family

VITAL SIGNS ON ARRIVAL:
- Blood Pressure: 92/54 mmHg
- Heart Rate: 112 bpm, irregularly irregular
- Temperature: 102.9 F
- Respiratory Rate: 26 breaths/min
- O2 Saturation: 93% on room air
- Glucose fingerstick: 148 mg/dL

ED IMPRESSION:
Suspected urosepsis with hypotension and altered mental status. Two large-bore IVs
placed, 30 mL/kg crystalloid initiated, blood and urine cultures drawn, stat labs
and lactate ordered. Empiric broad-spectrum antibiotics administered within 40
minutes of arrival per sepsis protocol. Urology and ICU consulted.
"""

LAB_BROOKS_PNG = """
Springfield Regional Laboratories
STAT LABORATORY REPORT

Patient: Harold Brooks
DOB: 04/30/1948 (Age 78)
MRN: MRN-2025-0112
Ordering Physician: Dr. Anita Rao, MD (Emergency Medicine)
Collection Date: 07/28/2026 08:35

CRITICAL VALUES REPORTED TO ORDERING PROVIDER AT 09:05

COMPLETE BLOOD COUNT (CBC):
- WBC: 19.8 x10^3/uL (H) [Ref: 4.5-11.0]
- Neutrophils: 91% (H) [Ref: 40-70]
- Bands: 8% - left shift present
- Hemoglobin: 11.8 g/dL (L) [Ref: 13.5-17.5]
- Platelets: 132 x10^3/uL (L) [Ref: 150-400]

CHEMISTRY:
- Lactate: 4.2 mmol/L CRITICAL (H) [Ref: 0.5-2.0]
- Creatinine: 2.3 mg/dL (H) [Baseline 1.6]
- BUN: 46 mg/dL (H) [Ref: 7-20]
- Glucose: 148 mg/dL (H) [Ref: 70-100]
- Sodium: 131 mEq/L (L) [Ref: 136-145]
- Potassium: 5.4 mEq/L (H) [Ref: 3.5-5.0]

URINALYSIS:
- Leukocyte esterase: LARGE
- Nitrite: POSITIVE
- WBC: Too numerous to count
- Bacteria: MANY

COAGULATION:
- INR: 1.4 (apixaban effect noted)
- Fibrinogen: 310 mg/dL (Normal)

BLOOD CULTURES x2: Collected, results pending

COMMENT: Findings consistent with sepsis secondary to urinary source with acute
kidney injury and lactate level meeting sepsis criteria. Serial lactates ordered.
"""

NOTES_BROOKS_DAY1 = """
St. Mary's Medical Center
PROGRESS NOTE - ICU DAY 1

Date: 07/29/2026
Patient: Harold Brooks
MRN: MRN-2025-0112
Attending: Dr. Priya Raman, Critical Care

SUBJECTIVE:
Patient admitted to MICU with urosepsis and septic shock. Overnight received 2.5L
crystalloid, norepinephrine started at 04:00 for refractory hypotension (MAP <60).
Daughter provides history; patient remains confused but follows simple commands.

OBJECTIVE:
Vital Signs:
- BP: 98/58 mmHg on norepinephrine 6 mcg/min
- HR: 96 bpm (afib rate-controlled)
- Temp: 100.4 F, downtrending
- RR: 20 breaths/min
- SpO2: 96% on 2L NC

Urine output: 45 mL/hr after Foley placement (retention relieved, drained 900 mL)

Repeat Labs (07/29 05:00):
- Lactate 2.8 (was 4.2) - clearing
- WBC 16.2 (was 19.8)
- Creatinine 2.5 - AKI peak not yet reached
- Blood cultures x2: Gram-positive cocci flagged at 14 hrs - STAT ID and
  sensitivities pending; vancomycin levels being followed

ASSESSMENT/PLAN:
1. Septic shock, urinary source - on sepsis bundle
   - Norepinephrine titration, wean as lactate clears
   - Piperacillin-tazobactam 3.375g IV q6h + vancomycin (load per protocol)
   - Escalation decision pending culture speciation
2. Acute kidney injury on CKD 3a - avoid nephrotoxins, hold ACEi
3. Afib - apixaban held 48h given AKI trend; cardiology aware
4. Delirium secondary to sepsis - reorientation, sleep protocol

Prognosis guarded. Family updated at bedside this morning.
"""

MED_LIST_BROOKS = """
St. Mary's Medical Center
DISCHARGE MEDICATION LIST

Patient: Harold Brooks
DOB: 04/30/1948
MRN: MRN-2025-0112
Prepared: 08/04/2026 by Pharmacy (K. Doyle, PharmD)

ACTIVE MEDICATIONS AT DISCHARGE:
1. Cephalexin 500 mg PO QID x 7 days (finish 08/10/2026)
   - Step-down oral therapy after IV course for ESBL E. coli urosepsis
2. Tamsulosin 0.4 mg PO daily (resume home medication)
3. Apixaban 5 mg PO BID (RESUMED 08/02 after AKI improved - renal dosing confirmed)
4. Atorvastatin 40 mg PO daily (resume home medication)

MEDICATIONS DISCONTINUED:
- Nitrofurantoin 100 mg PO BID suppressive therapy
   - STOPPED: inadequate suppression given recurrent resistant organisms;
     urology to discuss alternative prophylaxis at follow-up

ALLERGY ALERTS ON FILE:
- Codeine (hives) - avoid codeine and consider caution with other opioids

PHARMACIST COUNSELING:
- Complete cephalexin course even if symptoms resolve
- Return precautions reviewed: fever >101, confusion, minimal urination -> ED
"""

DISCHARGE_BROOKS = """
St. Mary's Medical Center
DISCHARGE SUMMARY

Patient: Harold Brooks
DOB: 04/30/1948 (Age 78)
MRN: MRN-2025-0112
Admission Date: 07/28/2026
Discharge Date: 08/04/2026
Attending Physician: Dr. Priya Raman, MD (Hospitalist service)

ADMISSION DIAGNOSIS:
1. Urosepsis with septic shock (A41.51)
2. Acute kidney injury on CKD stage 3a (N17.9)
3. Recurrent urinary tract infections
4. Benign prostatic hyperplasia
5. Atrial fibrillation

HOSPITAL COURSE:
78-year-old male admitted through the ED with fever, altered mental status, and
hypotension. Blood and urine cultures ultimately identified ESBL-producing E. coli
sensitive to piperacillin-tazobactam and cephalexin. Managed with IV fluids,
norepinephrine (weaned off hospital day 3), and targeted antibiotics. AKI peaked
at creatinine 2.6 on day 2 and returned to baseline 1.6 by discharge. Mental
status returned to baseline by day 4. Apixaban resumed after renal function
stabilized. Post-void bladder scans normal after tamsulosin continuation;
urology does not recommend catheterization.

CONDITION AT DISCHARGE:
Afebrile x 72 hours, hemodynamically stable, ambulating independently, mental
status at baseline.

DISCHARGE MEDICATIONS: See attached pharmacy medication list.

FOLLOW-UP:
- Infectious Disease clinic 08/11/2026
- Urology follow-up 08/20/2026 - discuss infection prophylaxis strategy
- Primary care 1 week post-discharge for renal panel recheck
- Home health nursing completed during admission stay transition

DISPOSITION:
Discharged home with daughter as caregiver. Home safety assessed.
"""


# ---------------------------------------------------------------------------
# Patient 5: Maria Delgado - 54F - community-acquired pneumonia (HIGH acuity)
# ---------------------------------------------------------------------------

INTAKE_DELGADO = """
Meadowbrook Community Hospital
PATIENT INTAKE FORM (EMERGENCY DEPARTMENT)

Date of Visit: 08/02/2026
Arrival Time: 15:40, walk-in

PATIENT INFORMATION:
First Name: Maria
Last Name: Delgado
Date of Birth: 02/09/1972 (Age 54)
Sex: Female
MRN: MRN-2025-0187
Phone: (555) 630-2251
Address: 340 Sycamore Blvd, Springfield, IL 62704
Insurance: BlueCross BlueShield PPO
Policy Number: BC-441-20987

EMERGENCY CONTACT:
Name: Carlos Delgado (Husband)
Phone: (555) 630-3390

CHIEF COMPLAINT:
Cough with fever for 4 days, right-sided chest pain

HISTORY OF PRESENT ILLNESS:
54-year-old female with 4 days of productive cough (rust-colored sputum), fevers
to 101.8 F, chills, and right-sided pleuritic chest pain. Shortness of breath when
climbing stairs, new this week. Tried OTC cold medicine without relief. Coworker
recently out with pneumonia.

PAST MEDICAL HISTORY:
1. Hypertension
2. Obesity
3. Seasonal allergic rhinitis

HOME MEDICATIONS:
1. Lisinopril 10 mg PO daily
2. Loratadine 10 mg PO PRN

ALLERGIES:
- No known drug allergies

FAMILY HISTORY:
- Mother: Breast cancer (survivor)
- Father: Hypertension

SOCIAL HISTORY:
- Occupation: Elementary school teacher
- Tobacco: Never smoker
- Alcohol: Occasional wine
- Vaccinations: Influenza this season - yes; pneumococcal - no

VITAL SIGNS ON ARRIVAL:
- Blood Pressure: 124/76 mmHg
- Heart Rate: 108 bpm
- Temperature: 101.6 F
- Respiratory Rate: 24 breaths/min
- O2 Saturation: 91% on room air

ED COURSE INITIALLY NOTED:
Chest x-ray performed: right lower lobe consolidation. Lab work ordered; IV placed.
Preliminary chest film read abnormal - radiology formal read pending. Admitted for
IV antibiotics under pneumonia severity index admission criteria.
"""

LAB_DELGADO = """
Meadowbrook Community Hospital
LABORATORY REPORT

Patient: Maria Delgado
DOB: 02/09/1972 (Age 54)
MRN: MRN-2025-0187
Ordering Physician: Dr. Kevin Alabi, MD
Collection Date: 08/02/2026 16:20

COMPLETE BLOOD COUNT (CBC):
- White Blood Cell Count: 15.2 x10^3/uL (HIGH) [Ref: 4.5-11.0]
- Neutrophils: 84% (HIGH) [Ref: 40-70]
- Lymphocytes: 9% (LOW) [Ref: 20-40]
- Hemoglobin: 13.1 g/dL (Normal) [Ref: 12.0-15.5]
- Platelet Count: 388 x10^3/uL (Normal) [Ref: 150-400]

INFLAMMATORY MARKERS:
- C-Reactive Protein: 182 mg/L (HIGH) [Ref: <10]
- Procalcitonin: 1.8 ng/mL (HIGH) [Ref: <0.5]

COMPREHENSIVE METABOLIC PANEL (CMP):
- Glucose: 104 mg/dL (Normal) [Ref: 70-100]
- Creatinine: 0.9 mg/dL (Normal) [Ref: 0.6-1.1]
- Sodium: 133 mEq/L (LOW) [Ref: 136-145]
- Remaining values within reference ranges

RESPIRATORY PATHOGEN PANEL (PCR): Negative for influenza A/B, SARS-CoV-2, RSV
SPUTUM GRAM STAIN: Many WBC, gram-positive diplococci seen
BLOOD CULTURES x2: Pending

CLINICAL SIGNIFICANCE:
Abnormal inflammatory markers with bacterial pneumonia pattern. Elevated
procalcitonin supports bacterial etiology. Hyponatremia common in lobar pneumonia.
Recommend empiric CAP coverage and repeat CRP to document treatment response.

Reported by: Microbiology & Core Lab
Verified by: Dr. Michael Torres, MD, Pathologist
Date: 08/02/2026
"""

NOTES_DELGADO_DAY1 = """
Meadowbrook Community Hospital
PROGRESS NOTE - HOSPITAL DAY 1

Date: 08/03/2026
Patient: Maria Delgado
MRN: MRN-2025-0187
Attending: Dr. Kevin Alabi, Internal Medicine

SUBJECTIVE:
Overnight received first doses ceftriaxone and azithromycin. Fevers continuing to
101.0 F. Cough persists, right-sided pain partially controlled with acetaminophen.
Sleeping with head elevated; still requiring supplemental oxygen.

OBJECTIVE:
Vital Signs:
- BP: 118/74 mmHg
- HR: 102 bpm
- Temp: 101.0 F
- RR: 22 breaths/min
- SpO2: 93% on 3L NC

Lungs: Decreased breath sounds right base with egophony; left clear
Remainder exam unremarkable

Radiology (formal): Right lower lobe consolidation with small parapneumonic
effusion - abnormal study consistent with lobar pneumonia.

Blood cultures: No growth at 24 hours.

ASSESSMENT/PLAN:
1. Community-acquired pneumonia, CURB-65 score 2 (RR >=22, O2 requirement)
   - Continue ceftriaxone 1g IV q24h + azithromycin 500mg IV q24h
   - Encourage incentive spirometry, mobilization
   - Titrate oxygen to SpO2 >=94%
2. Parapneumonic effusion, small - monitor; thoracic consult only if loculated
3. Hypertension - hold lisinopril while hypovolemic-febrile

Anticipate clinical improvement window 48-72 hours; will reassess daily.
"""

NOTES_DELGADO_DAY2 = """
Meadowbrook Community Hospital
PROGRESS NOTE - HOSPITAL DAY 2

Date: 08/04/2026
Patient: Maria Delgado
MRN: MRN-2025-0187
Attending: Dr. Kevin Alabi, Internal Medicine

SUBJECTIVE:
Slept better overnight. Reports feeling noticeably improved. Fever broke this
morning; last temp 99.2 F. Productive cough now clearing yellow sputum. Pain much
less with breathing.

OBJECTIVE:
Vital Signs:
- BP: 122/76 mmHg
- HR: 88 bpm
- Temp: 99.2 F
- RR: 18 breaths/min
- SpO2: 95% on 2L NC (weaning trial room air: 92%)

Labs: WBC 11.8 (down from 15.2), CRP 96 (down from 182) - appropriate response.
Blood cultures: No growth at 48 hours (final negative).

ASSESSMENT/PLAN:
1. CAP - responding to therapy; continue IV antibiotics today
   - Transition plan: oral levofloxacin 750 mg daily once afebrile x24h
2. Wean oxygen as tolerated; target room air by tomorrow morning
3. Ambulating in hall 3x today with nursing

Expect discharge eligibility tomorrow evening or next morning if trajectory holds.
"""

DISCHARGE_DELGADO = """
Meadowbrook Community Hospital
DISCHARGE SUMMARY

Patient: Maria Delgado
DOB: 02/09/1972 (Age 54)
MRN: MRN-2025-0187
Admission Date: 08/02/2026
Discharge Date: 08/05/2026
Attending Physician: Dr. Kevin Alabi, MD

ADMISSION DIAGNOSIS:
1. Community-acquired pneumonia, right lower lobe (J18.9)
2. Small parapneumonic effusion
3. Hypertension

HOSPITAL COURSE:
54-year-old female admitted with 4 days of productive cough, fever, and right-sided
pleuritic pain. Imaging showed RLL consolidation with small parapneumonic effusion.
Laboratory evaluation demonstrated leukocytosis with markedly elevated CRP and
procalcitonin. Treated with IV ceftriaxone and azithromycin, transitioned to oral
levofloxacin after becoming afebrile. Required up to 3L nasal cannula initially,
weaned to room air by hospital day 3. Effusion remained small and stable; no
drainage indicated. Blood cultures remained negative.

CONDITION AT DISCHARGE:
Afebrile x 36 hours, SpO2 96% on room air, ambulating without dyspnea at rest.

DISCHARGE MEDICATIONS:
1. Levofloxacin 750 mg PO daily x 5 more days (complete total course)
2. Lisinopril 10 mg PO daily (resumed)
3. Acetaminophen 650 mg PO q6h PRN pain

ALLERGIES: No known drug allergies

FOLLOW-UP:
- Primary care in 7 days with repeat chest x-ray to document resolution
- Repeat CRP not required if clinically improving
- Pneumococcal vaccination discussed - administer at PCP visit
- Return immediately for: fever recurrence, worsening dyspnea, hemoptysis

DISPOSITION:
Discharged to home. Activity as tolerated.
"""

# ---------------------------------------------------------------------------
# Patient 6: Tyler Nguyen - 29M - appendicitis, post-op (HIGH acuity)
# ---------------------------------------------------------------------------

LAB_NGUYEN = """
Meadowbrook Community Hospital
LABORATORY REPORT

Patient: Tyler Nguyen
DOB: 05/17/1997 (Age 29)
MRN: MRN-2025-0203
Ordering Physician: Dr. Rachel Kim, MD (General Surgery)
Collection Date: 08/12/2026 06:00 (Post-op Day 1)

COMPLETE BLOOD COUNT (CBC):
- White Blood Cell Count: 13.5 x10^3/uL (HIGH) [Ref: 4.5-11.0] - was 18.9 pre-op
- Neutrophils: 79% (HIGH) [Ref: 40-70]
- Hemoglobin: 14.6 g/dL (Normal) [Ref: 13.5-17.5]
- Platelet Count: 296 x10^3/uL (Normal) [Ref: 150-400]

COMPREHENSIVE METABOLIC PANEL (CMP):
- Glucose: 128 mg/dL (HIGH) [Ref: 70-100] - expected post-operative
- Creatinine: 0.95 mg/dL (Normal) [Ref: 0.7-1.3]
- Electrolytes: Within normal limits
- Liver panel: Within normal limits

INFLAMMATORY MARKER:
- C-Reactive Protein: 64 mg/L (HIGH) [Ref: <10] - down from 141 pre-op

URINALYSIS: Within normal limits

CLINICAL SIGNIFICANCE:
Trending improvement in leukocytosis and inflammatory markers following
laparoscopic appendectomy for non-perforated appendicitis. Values remain mildly
abnormal as expected on post-operative day 1. No evidence of bleeding or organ
dysfunction. Requires follow-up trending but no intervention indicated.

Reported by: Core Laboratory
Verified by: Dr. Michael Torres, MD, Pathologist
Date: 08/12/2026
"""

NOTES_NGUYEN_DOCX = """
Meadowbrook Community Hospital
PROGRESS NOTE - POST-OPERATIVE DAY 1

Date: 08/12/2026
Patient: Tyler Nguyen
MRN: MRN-2025-0203
Service: General Surgery (Dr. Rachel Kim)

SUBJECTIVE:
29-year-old male, post-op day 1 status post laparoscopic appendectomy for acute
non-perforated appendicitis. Slept reasonably well. States pain well controlled -
"maybe a 3 out of 10" at incision sites. Passed flatus this morning. Tolerating
clear liquids advanced to soft diet at lunch. Ambulated twice in hallway.

OBJECTIVE:
Vital Signs:
- BP: 122/74 mmHg
- HR: 76 bpm
- Temp: 99.0 F
- RR: 16 breaths/min
- SpO2: 99% room air

Abdomen: Benign, incisions clean dry intact, no erythema or drainage. Mild
tenderness at RLQ port site, no rebound. Bowel sounds present.

Labs: WBC 13.5 (down from 18.9), CRP 64 (down from 141) - appropriate post-op trend.

ASSESSMENT/PLAN:
1. Status post laparoscopic appendectomy, non-perforated - doing well
   - Advance diet as tolerated
   - Discontinue IV opioids, switch to ibuprofen 600mg q6h PRN + oxycodone 5mg PRN severe
   - Early ambulation encouraged
2. Anticipate discharge home tomorrow morning if tolerating diet and pain controlled
   orally.
Will review wound care instructions with patient this evening.
"""

DISCHARGE_NGUYEN = """
Meadowbrook Community Hospital
DISCHARGE SUMMARY

Patient: Tyler Nguyen
DOB: 05/17/1997 (Age 29)
MRN: MRN-2025-0203
Admission Date: 08/11/2026
Discharge Date: 08/13/2026
Attending Physician: Dr. Rachel Kim, MD (General Surgery)

ADMISSION DIAGNOSIS:
1. Acute non-perforated appendicitis (K35.80)

PROCEDURE PERFORMED:
- Laparoscopic appendectomy (08/11/2026, uncomplicated; pathology: acute
  appendicitis without perforation, margins benign)

HOSPITAL COURSE:
29-year-old healthy male who presented with 18 hours of periumbilical pain
migrating to the right lower quadrant with anorexia. CT demonstrated appendiceal
dilation with periappendiceal inflammation. Underwent uncomplicated laparoscopic
appendectomy the same evening. Post-operative course notable only for expected
mild leukocytosis, which trended down appropriately. Advanced diet POD1, ambulated
independently, pain controlled on oral analgesia alone by POD2.

CONDITION AT DISCHARGE:
Afebrile, tolerating regular diet, bowel function returning, comfortable on oral
analgesics.

DISCHARGE MEDICATIONS:
1. Ibuprofen 600 mg PO q6h PRN pain (with food)
2. Oxycodone 5 mg PO q4h PRN severe breakthrough pain (limited quantity dispensed)

ALLERGIES: No known drug allergies

WOUND CARE:
Incisions may be showered 24 hours post-op. Pat dry. No baths or swimming x 2 weeks.
Steri-strips will fall off on their own.

ACTIVITY:
No lifting >15 lbs x 2 weeks. May return to desk work in 1 week; sports in 4 weeks.

FOLLOW-UP:
- Surgery clinic in 10-14 days for incision check and pathology review
- Immediate return for: fever >101, wound redness/drainage, persistent vomiting,
  abdominal pain escalation

DISPOSITION: Discharged to home with brother present.
"""


# ---------------------------------------------------------------------------
# Patient 7: Grace Okafor - 41F - routine physical, new hypothyroidism (ROUTINE)
# ---------------------------------------------------------------------------

INTAKE_OKAFOR = """
Valley Health Clinic
ANNUAL PHYSICAL EXAM INTAKE FORM

Date of Visit: 07/14/2026

PATIENT INFORMATION:
First Name: Grace
Last Name: Okafor
Date of Birth: 09/28/1984 (Age 41)
Sex: Female
MRN: MRN-2025-0246
Phone: (555) 512-8844
Email: g.okafor@email.com
Address: 75 Birchwood Drive, Springfield, IL 62701
Insurance: United Healthcare
Policy Number: UHC-663-55102

EMERGENCY CONTACT:
Name: Chidi Okafor (Husband)
Phone: (555) 512-9921

REASON FOR VISIT:
Annual preventive physical examination

HISTORY OF PRESENT ILLNESS:
Generally well. Over past 3 months notes mild fatigue that persists despite adequate
sleep, slight cold intolerance, and 5 lb weight gain with unchanged diet. Skin feels
dry. Menstrual cycles regular. No hair loss patches, hoarseness, or constipation.

PAST MEDICAL HISTORY:
1. Seasonal allergies
2. Two prior uncomplicated pregnancies

SURGICAL HISTORY:
- Cesarean section (2016)
- Tubal ligation (2019)

CURRENT MEDICATIONS:
1. Cetirizine 10 mg PO daily (seasonal)
2. Prenatal vitamin PO daily

ALLERGIES:
- Sulfa drugs (rash)

FAMILY HISTORY:
- Mother: Hypothyroidism, diagnosed age 45
- Father: Colon cancer (diagnosed 62, survivor)

SOCIAL HISTORY:
- Occupation: Accountant
- Tobacco: Never smoker
- Alcohol: 1-2 drinks per month
- Exercise: Yoga 2x weekly
- Diet: Balanced, mostly vegetarian

SCREENING STATUS:
- Last pap smear: 01/2025 (normal)
- Mammogram: Not yet due (average risk)
- Colonoscopy: Not due (begin at age 52 given paternal history)

VITAL SIGNS:
- Blood Pressure: 116/74 mmHg
- Heart Rate: 62 bpm
- Temperature: 97.6 F
- Respiratory Rate: 14 breaths/min
- O2 Saturation: 99%
- Height: 5'4" (162.6 cm)
- Weight: 158 lbs (71.7 kg)
- BMI: 27.1 kg/m2

EXAM FINDINGS NOTED BY MA:
Neck supple, no visible thyromegaly. Skin dry bilateral shins. Otherwise
unremarkable intake observations documented for provider review.

PLAN (PROPOSED ORDERS PENDING PROVIDER):
1. Screening labs: CBC, CMP, lipid panel, TSH, HbA1c, vitamin D
2. Age-appropriate preventive counseling
"""

LAB_OKAFOR_PNG = """
Springfield Regional Laboratories
LABORATORY REPORT

Patient: Grace Okafor
DOB: 09/28/1984 (Age 41)
MRN: MRN-2025-0246
Ordering Physician: Dr. Sarah Williams, MD
Collection Date: 07/14/2026 09:30

THYROID FUNCTION:
- TSH: 8.9 mIU/L (H) [Ref: 0.4-4.0]
- Free T4: 0.8 ng/dL (L-normal border) [Ref: 0.8-1.8]

COMPLETE BLOOD COUNT (CBC):
- WBC: 6.1 x10^3/uL (Normal) [Ref: 4.5-11.0]
- Hemoglobin: 12.9 g/dL (Normal) [Ref: 12.0-15.5]
- Hematocrit: 39.1% (Normal) [Ref: 34.9-44.5]
- Platelets: 264 x10^3/uL (Normal) [Ref: 150-400]

COMPREHENSIVE METABOLIC PANEL (CMP):
All analytes within reference ranges
- Glucose: 89 mg/dL
- Creatinine: 0.78 mg/dL
- Sodium: 139 mEq/L
- Potassium: 4.1 mEq/L

LIPID PANEL:
- Total Cholesterol: 208 mg/dL (H) [Ref: <200]
- LDL: 121 mg/dL (H) [Ref: <100]
- HDL: 62 mg/dL (Normal)
- Triglycerides: 96 mg/dL (Normal)

OTHER:
- HbA1c: 5.2% (Normal) [Ref: <5.7]
- Vitamin D, 25-OH: 24 ng/mL (L) [Ref: 30-100]

COMMENT: Thyroid profile suggestive of subclinical hypothyroidism pattern with
borderline free T4. Lipid changes can accompany thyroid dysfunction. Consider
thyroid peroxidase antibody testing and provider review. Repeat TSH in 6 weeks if
initiating replacement therapy.
"""

NOTES_OKAFOR = """
Valley Health Clinic
PROGRESS NOTE - PREVENTIVE VISIT

Date: 07/14/2026
Patient: Grace Okafor
MRN: MRN-2025-0246
Provider: Dr. Sarah Williams, MD

SUBJECTIVE:
41-year-old female presenting for annual physical. Describes 3 months of mild
fatigue, cold intolerance, dry skin, and 5 lb weight gain. No depression symptoms
on screening. Review of systems otherwise negative.

OBJECTIVE:
Vital Signs: BP 116/74, HR 62, Temp 97.6, SpO2 99%
BMI 27.1

Physical Exam:
General: Alert, well-appearing
Neck: Supple, no thyromegaly or nodules palpable
Skin: Dry bilateral shins
Heart: Regular rhythm, no murmur
Lungs: Clear bilaterally
Abdomen: Soft, non-tender
Extremities: No edema; deep tendon reflexes with normal relaxation phase

Labs (same-day draw):
- TSH 8.9 mIU/L (above range), Free T4 borderline at 0.8
- Lipids: LDL 121, total cholesterol 208
- CBC, CMP, HbA1c: Normal; Vitamin D mildly insufficient at 24

ASSESSMENT:
1. Subclinical hypothyroidism pattern - likely autoimmune given maternal history;
   TPO antibodies pending
2. Borderline dyslipidemia, possibly thyroid-mediated
3. Vitamin D insufficiency

PLAN:
1. Start levothyroxine 50 mcg PO daily on empty stomach, 30 min before breakfast
2. TPO antibody lab added today
3. Repeat TSH/free T4 in 6 weeks; adjust dose to target
4. Vitamin D 2000 IU daily
5. Lifestyle counseling on lipid management; recheck lipids after thyroid
   optimization before considering statin discussion
6. Routine preventive screenings up to date; colonoscopy planning at age 52 given
   paternal history

Patient counseled on levothyroxine administration rules. Agreed with plan.
"""

# ---------------------------------------------------------------------------
# Patient 8: Samuel Patel - 63M - uncontrolled type 2 diabetes (HIGH acuity)
# ---------------------------------------------------------------------------

INTAKE_PATEL_DOCX = """
Valley Health Clinic
PATIENT INTAKE FORM

Date of Visit: 08/10/2026

PATIENT INFORMATION:
First Name: Samuel
Last Name: Patel
Date of Birth: 01/19/1963 (Age 63)
Sex: Male
MRN: MRN-2025-0298
Phone: (555) 887-1140
Email: sam.patel@email.com
Address: 210 Juniper Road, Springfield, IL 62702
Insurance: Medicare Advantage (Humana)
Policy Number: HUM-330-90217

EMERGENCY CONTACT:
Name: Anita Patel (Wife)
Phone: (555) 887-2288

CHIEF COMPLAINT:
Diabetes follow-up; very thirsty lately and sugars running high on glucometer

HISTORY OF PRESENT ILLNESS:
63-year-old male with 12-year history of type 2 diabetes returns after 8-month gap
in care. Home glucose logs show fasting values 190-260 most mornings. Reports
polyuria x2 nightly, constant thirst, and blurry vision by afternoon. Stopped
metformin 3 months ago due to cost after insurance change. No episodes of
hypoglycemia. No foot wounds; checks feet weekly.

PAST MEDICAL HISTORY:
1. Type 2 diabetes mellitus (2014)
2. Hypertension
3. Hyperlipidemia
4. Chronic kidney disease stage 2 (baseline creatinine 1.2)

HOME MEDICATIONS (REPORTED - PARTIALLY LAPSED):
1. Metformin 1000 mg PO BID - stopped by patient 3 months ago
2. Lisinopril 20 mg PO daily - taking irregularly
3. Atorvastatin 20 mg PO daily - taking irregularly

ALLERGIES:
- No known drug allergies

FAMILY HISTORY:
- Father: Type 2 Diabetes
- Sister: Type 2 Diabetes

SOCIAL HISTORY:
- Occupation: Retired postal worker
- Tobacco: Quit 2009 (15 pack-years)
- Alcohol: None
- Exercise: Walks 15 minutes some evenings
- Diet: Reports frequent rice and sweets portions larger than recommended

VITAL SIGNS:
- Blood Pressure: 154/90 mmHg
- Heart Rate: 78 bpm
- Temperature: 98.2 F
- Respiratory Rate: 16 breaths/min
- O2 Saturation: 98%
- Height: 5'9" (175.3 cm)
- Weight: 214 lbs (97.1 kg)
- BMI: 31.6 kg/m2

POINT-OF-CARE TESTING TODAY:
- Fingerstick glucose: 312 mg/dL
- Urine dipstick: trace protein

FOOT EXAM SCREENING (MA-documented):
Intact sensation monofilament plantar surfaces bilaterally, no ulcers, pulses 2+.

PROVIDER PLAN FLAGGED FROM INTAKE:
Same-day comprehensive metabolic panel and HbA1c ordered. Medication reconciliation
required - cost barrier addressed with $4 generic list counseling.
"""

LAB_PATEL = """
Quest Diagnostics
LABORATORY REPORT

Patient: Samuel Patel
DOB: 01/19/1963 (Age 63)
MRN: MRN-2025-0298
Ordering Physician: Dr. Sarah Williams, MD
Collection Date: 08/10/2026

DIABETES MONITORING:
- Hemoglobin A1c: 10.8% (HIGH) [Ref: <5.7] - Goal <7.0
- Fasting Glucose: 342 mg/dL (HIGH) [Ref: 70-100]

COMPREHENSIVE METABOLIC PANEL (CMP):
- Glucose: 342 mg/dL (HIGH) [Ref: 70-100]
- BUN: 24 mg/dL (HIGH) [Ref: 7-20]
- Creatinine: 1.4 mg/dL (HIGH) [Ref: 0.7-1.3] - Baseline 1.2
- eGFR: 51 mL/min/1.73m2 (LOW) [Ref: >60]
- Sodium: 137 mEq/L (Normal) [Ref: 136-145]
- Potassium: 4.6 mEq/L (Normal) [Ref: 3.5-5.0]
- CO2: 25 mEq/L (Normal) [Ref: 23-29]
- Calcium: 9.4 mg/dL (Normal) [Ref: 8.5-10.5]
- Albumin: 4.2 g/dL (Normal) [Ref: 3.5-5.0]
- ALT: 34 U/L (Normal) [Ref: 7-56]
- AST: 29 U/L (Normal) [Ref: 10-40]
- Alkaline Phosphatase: 82 U/L (Normal) [Ref: 44-147]

URINE MICROALBUMIN/CREATININE RATIO: 118 mcg/mg (HIGH) [Ref: <30]

LIPID PANEL:
- Total Cholesterol: 218 mg/dL (HIGH) [Ref: <200]
- LDL Cholesterol: 138 mg/dL (HIGH) [Ref: <100 - diabetic goal]
- HDL Cholesterol: 34 mg/dL (LOW) [Ref: >40]
- Triglycerides: 232 mg/dL (HIGH) [Ref: <150]

CLINICAL SIGNIFICANCE:
Markedly elevated A1c and fasting glucose indicate severely uncontrolled diabetes.
Elevated urine microalbumin with declining eGFR suggests early diabetic nephropathy
progression. Dyslipidemia below diabetic targets. Recommend prompt medication
re-initiation with renal-adjusted dosing consideration and insulin initiation
evaluation. Ketones were not detected on urinalysis; no acidosis pattern present.

Reported by: Quest Diagnostics Core Lab
Verified by: Dr. Michael Torres, MD, Pathologist
Date: 08/10/2026
"""

NOTES_PATEL = """
Valley Health Clinic
PROGRESS NOTE

Date: 08/10/2026
Patient: Samuel Patel
MRN: MRN-2025-0298
Provider: Dr. Sarah Williams, MD

SUBJECTIVE:
63-year-old male with poorly controlled type 2 diabetes presenting after extended
lapse in care. Confirms stopping metformin due to cost barrier and inconsistent
use of antihypertensives. Polydipsia and nocturnal polyuria prominent. Vision
blurry late in day, worse over past 6 weeks. No chest pain, no foot wounds,
no hypoglycemic events.

OBJECTIVE:
Vital Signs: BP 154/90, HR 78, SpO2 98%, Weight 214 lbs, BMI 31.6
Point-of-care glucose: 312 mg/dL
Monofilament foot exam: intact sensation, no deformity, pulses 2+

Labs reviewed (drawn today):
- A1c 10.8%, fasting glucose 342
- eGFR 51 (down from baseline), urine ACR 118
- LDL 138, HDL 34, triglycerides 232

ASSESSMENT:
1. Type 2 diabetes mellitus, severely uncontrolled (A1c 10.8%) with early diabetic
   nephropathy
2. Hypertension, above goal
3. Dyslipidemia, below target
4. CKD stage 3a - new staging based on today's eGFR
5. Medication non-adherence driven by cost - social determinant barrier identified

PLAN:
1. Restart metformin 500 mg PO BID (renally acceptable at eGFR 51; do not titrate
   higher without nephrology input)
2. Initiate basal insulin glargine 10 units SC nightly; teach injection technique
   today with nurse educator; hypoglycemia precautions reviewed
3. Add empagliflozin 10 mg PO daily - renal and cardioprotective benefit
4. Resume lisinopril 20 mg daily and atorvastatin 20 mg nightly; enrolled in $4
   generic program to remove cost barrier
5. Diabetic retinopathy screening - ophthalmology referral placed (blurry vision)
6. Nutrition referral; carbohydrate portion counseling emphasized
7. Repeat BMP in 2 weeks to reassess renal function on new agents
8. Return sooner if glucose >400, symptoms of dehydration, or any foot wound

Patient motivated after reviewing numbers together. Wife included in education.
Follow-up appointment scheduled 08/24/2026.
"""


# ---------------------------------------------------------------------------
# Patient 9: Ava Simmons - 8F - asthma exacerbation (ROUTINE/MILD acuity)
# ---------------------------------------------------------------------------

INTAKE_SIMMONS = """
Pediatric Care Associates of Springfield
PATIENT INTAKE FORM (PEDIATRIC)

Date of Visit: 08/18/2026

PATIENT INFORMATION:
First Name: Ava
Last Name: Simmons
Date of Birth: 06/02/2018 (Age 8)
Sex: Female
MRN: MRN-2025-0334
Guardian: Megan Simmons (Mother)
Guardian Phone: (555) 905-4417
Address: 61 Dogwood Circle, Springfield, IL 62705
Insurance: Medicaid IL
Policy Number: ILD-772-44903

EMERGENCY CONTACT:
Name: Jordan Simmons (Father)
Phone: (555) 905-5528

CHIEF COMPLAINT:
Wheezing and nighttime cough for 4 days after summer cold

HISTORY OF PRESENT ILLNESS (per guardian):
8-year-old female with known mild persistent asthma presents after 4 days of
wheezing audible at night and cough disrupting sleep 2-3 times per night. Used
albuterol inhaler 3 times yesterday with partial relief lasting about 3 hours each.
Running nose improved but wheeze persists. Missing swim practice this week. No
fever. No prior ICU admissions; one urgent care visit for asthma last spring.

PAST MEDICAL HISTORY:
1. Mild persistent asthma (diagnosed age 5)
2. Seasonal allergic rhinitis
3. Peanut allergy (anaphylaxis history - epinephrine autoinjector prescribed)

ASTHMA CONTROLLER STATUS:
- Was prescribed fluticasone 44 mcg 2 puffs BID - guardian reports stopped using
  in June when child "seemed fine"

ALLERGIES:
- PEANUT - ANAPHYLAXIS (carries epinephrine autoinjector)
- Dust mites (environmental)

CURRENT MEDICATIONS:
1. Albuterol HFA 90 mcg 2 puffs q4-6h PRN (increased use this week per guardian)
2. Cetirizine 5 mg PO daily
3. Epinephrine autoinjector 0.15 mg IM PRN anaphylaxis

FAMILY HISTORY:
- Mother: Asthma, seasonal allergies
- Maternal uncle: Asthma

SOCIAL HISTORY:
- Lives with parents and older brother; cat in home
- School: Lincoln Elementary (school nurse has rescue inhaler on file)
- Exposures: Swimming 2x/week (indoor pool)

VITAL SIGNS (PEDIATRIC):
- Blood Pressure: 102/64 mmHg
- Heart Rate: 110 bpm
- Temperature: 98.9 F
- Respiratory Rate: 24 breaths/min
- O2 Saturation: 95% on room air
- Weight: 57 lbs (25.9 kg)

PRE-PROVIDER TRIAGE NOTES:
No audible wheeze; expiratory wheezes noted on auscultation; speaking in full
sentences. Peak flow performed: 65% of personal best. Nebulized albuterol with
ipratropium treatment started per standing order pending provider evaluation.
"""

LAB_SIMMONS_PNG = """
Springfield Regional Laboratories
PEDIATRIC LABORATORY REPORT

Patient: Ava Simmons
DOB: 06/02/2018 (Age 8)
MRN: MRN-2025-0334
Ordering Physician: Dr. Miguel Santos, MD (Pediatrics)
Collection Date: 08/18/2026

COMPLETE BLOOD COUNT WITH DIFFERENTIAL:
- WBC: 7.8 x10^3/uL (Normal) [Ref: 5.0-13.0]
- Hemoglobin: 13.2 g/dL (Normal) [Ref: 11.5-15.5]
- Hematocrit: 39.4% (Normal) [Ref: 34.0-44.0]
- Platelets: 342 x10^3/uL (Normal) [Ref: 150-450]
- Absolute Eosinophil Count: 520 cells/uL (H) [Ref: 30-350]
- Total IgE: 148 kU/L (H) [Ref: <100 for age]

ALLERGY TESTING (SERUM SPECIFIC IgE):
- D. farinae (dust mite): Positive Class 3
- Cat dander: Positive Class 2
- Peanut: Positive Class 4 - CONSISTENT WITH KNOWN ANAPHYLAXIS HISTORY
- Alternaria mold: Positive Class 2

PULSE OXIMETRY (ROOM AIR, POST-TREATMENT): 97%

COMMENT: Eosinophilia and raised IgE support allergic asthma phenotype. Multiple
environmental sensitizations identified - dust mite mitigation counseling
recommended along with pet exposure discussion. Viral studies not indicated given
improving upper respiratory symptoms and afebrile status. No acute infection
identified.
"""

NOTES_SIMMONS = """
Pediatric Care Associates of Springfield
PROGRESS NOTE

Date: 08/18/2026
Patient: Ava Simmons
MRN: MRN-2025-0334
Provider: Dr. Miguel Santos, MD (Pediatrics)
Guardian Present: Megan Simmons (Mother)

SUBJECTIVE:
8-year-old with mild persistent asthma, non-adherent to controller since June,
presenting after viral URI with 4 nights of wheeze and nocturnal cough. Albuterol
providing only 3-hour relief intervals. Received one nebulized albuterol with
ipratropium treatment at check-in with good response. No fever. Peanut allergy
with autoinjector - guardian confirms device current and accessible.

OBJECTIVE:
Vitals: HR 96 (post-treatment), RR 20, SpO2 97% room air, Temp 98.9 F
Peak flow repeat: 82% of personal best (up from 65% pre-treatment)

Exam:
General: Interactive, playful, speaking full sentences
Lungs: Improved aeration, faint expiratory wheeze right mid-field, no retractions,
accessory muscle use absent
ENT/Nose: Minimal clear rhinorrhea, turbinates pale
Rest of exam: Unremarkable

Labs: CBC with eosinophils 520 (allergic phenotype marker), specific IgE panel
positive for dust mite/cat/Alternaria; peanut Class 4 confirming known allergy.

ASSESSMENT:
1. Acute mild-moderate asthma exacerbation on background of controller
   non-adherence, triggered by viral URI plus environmental exposures
2. Allergic rhinitis with multiple environmental sensitivities
3. Peanut anaphylaxis history - action plan verified current

PLAN:
1. Resume fluticasone 44 mcg 2 puffs BID with spacer - adherence importance
   discussed extensively with guardian and child (child chose sticker chart system)
2. Albuterol PRN q4h for next 48 hours, then PRN only; return if needing it more
   than 2x/day after 48h
3. Oral prednisolone 15 mg daily x 3 days burst provided
4. Asthma action plan updated and printed - green/yellow/red zones reviewed;
   copy for school nurse faxed
5. Environmental controls: allergen-proof pillow/mattress covers, HEPA filter in
   bedroom, honest discussion of cat exposure impact with family
6. Allergy/immunology referral placed for sensitization management
7. Follow-up in 2 weeks to confirm controller adherence and symptom scores

Guardian verbalized understanding of all medication changes. Child cooperative
throughout visit. Return precautions clearly stated.
"""

# ---------------------------------------------------------------------------
# Patient 10: Dorothy Whitfield - 71F - COPD w/ respiratory failure (CRITICAL)
# ---------------------------------------------------------------------------

LAB_WHITFIELD = """
Meadowbrook Community Hospital
LABORATORY REPORT

Patient: Dorothy Whitfield
DOB: 11/07/1954 (Age 71)
MRN: MRN-2025-0391
Ordering Physician: Dr. Elena Vasquez, MD (Pulmonology)
Collection Date: 08/15/2026 22:40

ARTERIAL BLOOD GAS (ROOM AIR AT TRIAGE):
- pH: 7.21 CRITICAL LOW [Ref: 7.35-7.45]
- pCO2: 88 mmHg CRITICAL HIGH [Ref: 35-45]
- pO2: 52 mmHg (LOW) [Ref: 80-100]
- HCO3: 34 mEq/L (HIGH) [Ref: 22-28] - chronic retention component
- Lactate: 2.4 mmol/L (HIGH) [Ref: 0.5-2.0]

CRITICAL VALUE CALLED TO ED PHYSICIAN AT 23:05

COMPLETE BLOOD COUNT (CBC):
- White Blood Cell Count: 12.8 x10^3/uL (HIGH) [Ref: 4.5-11.0]
- Neutrophils: 81% (HIGH)
- Hemoglobin: 16.8 g/dL (HIGH) [Ref: 12.0-15.5] - chronic hypoxemia compensation
- Hematocrit: 50.1% (HIGH)
- Platelet Count: 301 x10^3/uL (Normal)

CHEMISTRY:
- Potassium: 3.2 mEq/L (LOW) [Ref: 3.5-5.0]
- Magnesium: 1.6 mg/dL (LOW) [Ref: 1.7-2.2]
- Creatinine: 1.0 mg/dL (Normal)
- Glucose: 156 mg/dL (HIGH) - steroid effect from home prednisone burst

THEOPHYLLINE LEVEL: 6.2 mcg/mL [Therapeutic: 5-15]

CLINICAL SIGNIFICANCE:
Critical blood gas demonstrating acute-on-chronic hypercapnic respiratory failure
with severe acidemia in a patient with known severe COPD. Electrolyte deficits
require repletion. Findings prompted urgent BiPAP initiation and ICU evaluation.

Reported by: Critical Care Laboratory
Verified by: Dr. Michael Torres, MD, Pathologist
Date: 08/15/2026
"""

NOTES_WHITFIELD_DAY1 = """
Meadowbrook Community Hospital
PROGRESS NOTE - MICU DAY 1

Date: 08/16/2026
Patient: Dorothy Whitfield
MRN: MRN-2025-0391
Attending: Dr. Elena Vasquez, Pulmonology/Critical Care

SUBJECTIVE:
71-year-old female with severe COPD (GOLD stage III, home O2 1L nocturnal) brought
by EMS after 3 days of worsening dyspnea and productive green sputum; found by
neighbor somnolent this evening. In ED, developed CO2 narcosis pattern - drowsy,
asterixis present. Started on BiPAP 12/6 with rapid mental status improvement.
Family reports she ran out of her tiotropium 5 weeks ago and did not refill.

OBJECTIVE:
Vital Signs:
- BP: 138/82 mmHg
- HR: 108 bpm
- Temp: 100.8 F
- RR: 26 -> 20 breaths/min on BiPAP
- SpO2: 88% -> 94% on BiPAP, FiO2 40%

Repeat ABG on BiPAP (01:30): pH 7.29, pCO2 68, pO2 68 - improving
Chest x-ray: hyperinflation with right lower lobe infiltrate - abnormal
Sputum culture: sent; empiric coverage started

ASSESSMENT/PLAN:
1. Acute on chronic hypercapnic respiratory failure secondary to COPD exacerbation
   with community-acquired pneumonia - admitted to MICU
   - BiPAP 12/6, titrate FiO2 targeting SpO2 88-92% (permissive hypercapnia)
   - Methylprednisolone 40mg IV q12h
   - Ceftriaxone + azithromycin for CAP coverage
   - Duoneb q4h scheduled
   - If mental status declines or ABG worsens despite NIV -> intubation discussion;
     patient has advance directive on file - family meeting planned today
2. Hypokalemia/hypomagnesemia - repletion protocols started
3. Theophylline therapeutic - continue home dose

Patient more alert this morning, following commands. Respiratory therapist notes
good mask tolerance.
"""

NOTES_WHITFIELD_DAY2 = """
Meadowbrook Community Hospital
PROGRESS NOTE - MICU DAY 2

Date: 08/17/2026
Patient: Dorothy Whitfield
MRN: MRN-2025-0391
Attending: Dr. Elena Vasquez, Pulmonology/Critical Care

SUBJECTIVE:
Much improved overnight. Off BiPAP for morning trial with toleration. Speaking in
short sentences, states she "can breathe again." Afebrile since midnight. Family
meeting held - goals of care reviewed; patient wishes full support short of
prolonged ventilation, consistent with documented preferences.

OBJECTIVE:
Vital Signs:
- BP: 126/76 mmHg
- HR: 92 bpm
- Temp: 98.9 F
- RR: 18 breaths/min
- SpO2: 92% on 2L NC

ABG on 2L NC (06:00): pH 7.34, pCO2 58, pO2 62 - compensated, acceptable baseline+
Labs: WBC 10.9 (down from 12.8), lactate cleared to 1.4, K+ repleted to 4.0
Sputum culture: H. influenzae, pan-sensitive - de-escalated to amoxicillin-clavulanate

ASSESSMENT/PLAN:
1. Hypercapnic respiratory failure - resolving on non-invasive support
   - Remain off BiPAP if ABG stable on repeat; high-flow nasal cannula bridge PRN
   - Continue steroid taper plan: IV today, transition to prednisone 40mg PO x5d
2. CAP due to H. influenzae - complete 7-day amoxicillin-clavulanate course
3. COPD maintenance reconstruction: tiotropium restarted today (lapse addressed -
   will enroll in medication delivery program before discharge)
4. Transfer to medical floor this afternoon if stable; anticipate discharge in
   3-4 days with home oxygen re-evaluation

Respiratory therapy weaning protocol in effect.
"""

DISCHARGE_WHITFIELD = """
Meadowbrook Community Hospital
DISCHARGE SUMMARY

Patient: Dorothy Whitfield
DOB: 11/07/1954 (Age 71)
MRN: MRN-2025-0391
Admission Date: 08/15/2026
Discharge Date: 08/21/2026
Attending Physician: Dr. Elena Vasquez, MD (Pulmonology)

ADMISSION DIAGNOSIS:
1. Acute on chronic hypercapnic respiratory failure (J96.11)
2. COPD exacerbation with community-acquired pneumonia (J44.1)
3. Severe COPD, GOLD stage III
4. Medication lapse - tiotropium non-refill (contributory factor)

HOSPITAL COURSE:
71-year-old female with severe COPD admitted via ED with somnolence and critical
hypercapnia (pH 7.21, pCO2 88). Managed successfully with BiPAP, avoiding
intubation, consistent with her documented goals of care. Treated for COPD
exacerbation with corticosteroids and bronchodilators; pneumonia culture-positive
for H. influenzae, treated with targeted oral antibiotics after initial IV
empiric therapy. Transferred to floor day 2, ambulating with supplemental oxygen
by day 3. Home oxygen qualification testing completed - qualifies for 24-hour
supplemental oxygen at discharge setting of 2L.

CONDITION AT DISCHARGE:
Stable on room air exercises with exertional desaturation managed by prescribed
oxygen. No dyspnea at rest. Ambulating 150 feet independently.

DISCHARGE MEDICATIONS:
1. Prednisone 40 mg PO daily x 5 days (completes taper)
2. Amoxicillin-clavulanate 875/125 mg PO BID x 7 total days (day 4 of course)
3. Tiotropium 18 mcg inhaled daily (RESTARTED - critical maintenance medication)
4. Budesonide/formoterol inhaler 2 puffs BID
5. Albuterol HFA PRN q4h
6. Theophylline 300 mg PO daily (level checked at follow-up)

ALLERGIES: No known drug allergies

OXYGEN ORDERS:
- Oxygen 2 L/min continuous via nasal cannula; 4 L/min with exertion
- DME supplier arranged; conserving device provided

FOLLOW-UP:
- Pulmonology clinic 08/28/2026 - theophylline level, ABG reassessment
- Home health respiratory nursing visits 3x/week x 2 weeks
- Pulmonary rehabilitation referral submitted
- Smoking cessation counseling continued (quit 2019 - relapse prevention)

DISPOSITION:
Discharged to home; neighbor support network confirmed with patient consent.
Return precautions including increased somnolence or sputum change reviewed.
"""


# ---------------------------------------------------------------------------
# Build manifest
# ---------------------------------------------------------------------------

JOBS = [
    # Patient 1: Margaret Thompson (MRN-2024-0847) - CHF/T2DM/HTN
    ("pdf", INTAKE_THOMPSON, "intake_form_thompson.pdf"),
    ("pdf", LAB_THOMPSON, "lab_report_thompson.pdf"),
    # Patient 2: Robert Martinez (MRN-2024-1203) - new-onset T2DM
    ("pdf", INTAKE_MARTINEZ, "intake_form_martinez.pdf"),
    ("txt", NOTES_MARTINEZ, "physician_notes_martinez.txt"),
    # Patient 3: Emily Johnson (MRN-2024-1567) - fatigue/headaches
    ("pdf", INTAKE_JOHNSON, "intake_form_johnson.pdf"),
    ("txt", FOLLOWUP_NOTE_JOHNSON, "physician_notes_johnson_followup.txt"),
    # Patient 4: Harold Brooks (MRN-2025-0112) - urosepsis [CRITICAL]
    ("pdf", INTAKE_BROOKS, "intake_form_brooks.pdf"),
    ("png", LAB_BROOKS_PNG, "lab_report_brooks.png"),
    ("txt", NOTES_BROOKS_DAY1, "physician_notes_brooks_day1.txt"),
    ("txt", MED_LIST_BROOKS, "medication_list_brooks.txt"),
    ("docx", DISCHARGE_BROOKS, "discharge_summary_brooks.docx"),
    # Patient 5: Maria Delgado (MRN-2025-0187) - pneumonia [HIGH]
    ("pdf", INTAKE_DELGADO, "intake_form_delgado.pdf"),
    ("pdf", LAB_DELGADO, "lab_report_delgado.pdf"),
    ("txt", NOTES_DELGADO_DAY1, "physician_notes_delgado_day1.txt"),
    ("txt", NOTES_DELGADO_DAY2, "physician_notes_delgado_day2.txt"),
    ("pdf", DISCHARGE_DELGADO, "discharge_summary_delgado.pdf"),
    # Patient 6: Tyler Nguyen (MRN-2025-0203) - appendicitis post-op [HIGH]
    ("pdf", LAB_NGUYEN, "lab_report_nguyen_postop.pdf"),
    ("docx", NOTES_NGUYEN_DOCX, "physician_notes_nguyen.docx"),
    ("pdf", DISCHARGE_NGUYEN, "discharge_summary_nguyen.pdf"),
    # Patient 7: Grace Okafor (MRN-2025-0246) - routine physical [ROUTINE]
    ("pdf", INTAKE_OKAFOR, "intake_form_okafor.pdf"),
    ("png", LAB_OKAFOR_PNG, "lab_report_okafor.png"),
    ("txt", NOTES_OKAFOR, "physician_notes_okafor.txt"),
    # Patient 8: Samuel Patel (MRN-2025-0298) - uncontrolled T2DM [HIGH]
    ("docx", INTAKE_PATEL_DOCX, "intake_form_patel.docx"),
    ("pdf", LAB_PATEL, "lab_report_patel.pdf"),
    ("txt", NOTES_PATEL, "physician_notes_patel.txt"),
    # Patient 9: Ava Simmons (MRN-2025-0334) - pediatric asthma [ROUTINE]
    ("pdf", INTAKE_SIMMONS, "intake_form_simmons.pdf"),
    ("png", LAB_SIMMONS_PNG, "lab_report_simmons.png"),
    ("txt", NOTES_SIMMONS, "physician_notes_simmons.txt"),
    # Patient 10: Dorothy Whitfield (MRN-2025-0391) - COPD resp failure [CRITICAL]
    ("pdf", LAB_WHITFIELD, "lab_report_whitfield_abg.pdf"),
    ("txt", NOTES_WHITFIELD_DAY1, "physician_notes_whitfield_day1.txt"),
    ("txt", NOTES_WHITFIELD_DAY2, "physician_notes_whitfield_day2.txt"),
    ("pdf", DISCHARGE_WHITFIELD, "discharge_summary_whitfield.pdf"),
]

WRITERS = {"pdf": create_pdf, "docx": create_docx, "txt": create_txt, "png": create_png}


def main():
    print("Generating synthetic patient-document dataset...")
    for fmt, content, filename in JOBS:
        WRITERS[fmt](content, filename)
    print(f"\nDone! {len(JOBS)} documents written to {SAMPLES_DIR}")


if __name__ == "__main__":
    main()




